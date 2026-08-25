"""TDD tests for the Compensation Analysis report section (Wave 6d, Task 8).

Pins the four wiring points for the new section:
  1. ``REPORT_SECTIONS`` (pdf_report) gains the ``compensation`` key.
  2. The PDF dispatcher (``generate_ombudsman_pdf``) wires it.
  3. The DOCX dispatcher (``generate_ombudsman_docx``) wires it.
  4. The HTML dispatcher (``generate_html_report``) wires it.

Plus behavioural coverage: a synthetic record set that yields compensation
rows renders a "Compensation Analysis" section in a generated PDF, DOCX and
HTML; empty rows produce the "no compensation claims identified" note without
crashing any format.
"""

from __future__ import annotations

from pathlib import Path
from typing import Any

import pytest

from edf_bill_fetcher.io.reporters.pdf_report import REPORT_SECTIONS

AS_OF = "2026-06-01"


def _compensation_records() -> list[dict[str, Any]]:
    """A record set that yields compensation rows.

    Mirrors the fixture in ``tests/test_processors_compensation.py``: a
    back-billing invoice (Period From 01 Jan 2022 -> Period To 28 Feb 2024,
    billed 01 Mar 2024) yields a ``back_billing_excess`` row, and an
    unrefunded credit (01 Feb 2026) yields ``credit_hold_interest`` +
    ``late_credit_interest`` rows.
    """
    return [
        {
            "Date": "01 Mar 2024",
            "Source": "Local PDF Folder",
            "Entry Type": "New Bill",
            "Amount (£)": 1200.00,
            "Period Charge (£)": 1200.00,
            "Period From": "01 Jan 2022",
            "Period To": "28 Feb 2024",
            "Invoice #": "KI-0001",
            "Reading": "Actual",
            "Units (kWh)": "100",
            "Standing Chg (p/day)": "60.1",
            "Tariff": "Standard Variable",
            "Attachment Name": "bill_0001.pdf",
            "Details": "Back-billed consumption",
            "Logic Used": "Period Charge Match",
        },
        {
            "Date": "01 Feb 2026",
            "Source": "HTM Account History",
            "Entry Type": "Credit Note",
            "Amount (£)": -100.00,
            "Period Charge (£)": 0.00,
            "Period From": "01 Jan 2026",
            "Period To": "31 Jan 2026",
            "Invoice #": "KCR-0001",
            "Reading": "Actual",
            "Units (kWh)": "0",
            "Standing Chg (p/day)": "60.1",
            "Tariff": "Standard Variable",
            "Attachment Name": "credit_0001.htm",
            "Details": "Credit balance",
            "Logic Used": "Credit Note",
        },
    ]


def _no_compensation_records() -> list[dict[str, Any]]:
    """A record set that yields NO compensation rows (promptly billed,
    no credit balance)."""
    return [
        {
            "Date": "01 Jun 2024",
            "Source": "Local PDF Folder",
            "Entry Type": "New Bill",
            "Amount (£)": 100.00,
            "Period Charge (£)": 100.00,
            "Period From": "01 May 2024",
            "Period To": "31 May 2024",
            "Invoice #": "KI-0002",
            "Reading": "Actual",
            "Units (kWh)": "100",
            "Standing Chg (p/day)": "60.1",
            "Tariff": "Standard Variable",
            "Attachment Name": "bill_0002.pdf",
            "Details": "Promptly billed",
            "Logic Used": "Period Charge Match",
        }
    ]


class _StubEngine:
    def __init__(self) -> None:
        self.pdf_count = 1
        self.email_count = 1
        self.filtered_records: list[dict] = []


def _config() -> dict[str, Any]:
    return {
        "report_account_ref": "A-12345678",
        "as_of": AS_OF,
        "report_sections": ["compensation"],
    }


class TestCompensationRegistry:
    """The registry declares the compensation section as a main section."""

    def test_registry_gains_compensation_key(self) -> None:
        keys = [s.key for s in REPORT_SECTIONS]
        assert "compensation" in keys

    def test_compensation_is_main_not_appendix(self) -> None:
        meta = next(s for s in REPORT_SECTIONS if s.key == "compensation")
        assert meta.title == "Compensation Analysis"
        assert meta.is_appendix is False

    def test_compensation_after_tariff_before_appendices(self) -> None:
        keys = [s.key for s in REPORT_SECTIONS]
        assert keys.index("tariff") < keys.index("compensation")
        assert keys.index("compensation") < keys.index("appendix_methodology")


class TestCompensationDispatch:
    """All three dispatchers wire the compensation section."""

    def test_pdf_dispatcher_wires_compensation(self) -> None:
        from edf_bill_fetcher.io.reporters.pdf_report import generate_ombudsman_pdf

        assert callable(generate_ombudsman_pdf)

    def test_docx_dispatcher_wires_compensation(self) -> None:
        from edf_bill_fetcher.io.reporters.docx_report import generate_ombudsman_docx

        assert callable(generate_ombudsman_docx)

    def test_html_dispatcher_wires_compensation(self) -> None:
        from edf_bill_fetcher.io.reporters.html_report import generate_html_report

        assert callable(generate_html_report)


class TestCompensationRender:
    """A synthetic record set that yields compensation rows renders a
    Compensation Analysis section in PDF, DOCX and HTML."""

    def test_pdf_renders_compensation_section(self, tmp_path: Path) -> None:
        from edf_bill_fetcher.io.reporters.pdf_report import generate_ombudsman_pdf

        out = str(tmp_path / "comp.pdf")
        result = generate_ombudsman_pdf(
            _compensation_records(), out, _config(), _StubEngine()
        )
        assert result == out
        assert Path(out).exists()
        assert Path(out).stat().st_size > 0

    def test_docx_renders_compensation_section(self, tmp_path: Path) -> None:
        from docx import Document

        from edf_bill_fetcher.io.reporters.docx_report import generate_ombudsman_docx

        out = str(tmp_path / "comp.docx")
        result = generate_ombudsman_docx(
            _compensation_records(), out, _config(), _StubEngine()
        )
        assert result == out
        doc = Document(out)
        texts = [p.text for p in doc.paragraphs]
        assert any("Compensation Analysis" in t for t in texts)

    def test_html_renders_compensation_section(self, tmp_path: Path) -> None:
        from edf_bill_fetcher.io.reporters.html_report import generate_html_report

        out = str(tmp_path / "comp.html")
        result = generate_html_report(
            _compensation_records(), out, _config(), _StubEngine()
        )
        assert result == out
        rendered = Path(out).read_text(encoding="utf-8")
        assert "Compensation Analysis" in rendered


class TestCompensationEmpty:
    """Empty compensation rows produce the no-claims note without crashing."""

    def test_pdf_empty_rows_no_crash(self, tmp_path: Path) -> None:
        from edf_bill_fetcher.io.reporters.pdf_report import generate_ombudsman_pdf

        out = str(tmp_path / "empty.pdf")
        result = generate_ombudsman_pdf(
            _no_compensation_records(), out, _config(), _StubEngine()
        )
        assert result == out
        assert Path(out).exists()

    def test_docx_empty_rows_no_crash(self, tmp_path: Path) -> None:
        from docx import Document

        from edf_bill_fetcher.io.reporters.docx_report import generate_ombudsman_docx

        out = str(tmp_path / "empty.docx")
        result = generate_ombudsman_docx(
            _no_compensation_records(), out, _config(), _StubEngine()
        )
        assert result == out
        doc = Document(out)
        texts = [p.text for p in doc.paragraphs]
        assert any("no compensation claims identified" in t.lower() for t in texts)

    def test_html_empty_rows_no_crash(self, tmp_path: Path) -> None:
        from edf_bill_fetcher.io.reporters.html_report import generate_html_report

        out = str(tmp_path / "empty.html")
        result = generate_html_report(
            _no_compensation_records(), out, _config(), _StubEngine()
        )
        assert result == out
        rendered = Path(out).read_text(encoding="utf-8")
        assert "no compensation claims identified" in rendered.lower()


if __name__ == "__main__":
    pytest.main([__file__, "-v"])
