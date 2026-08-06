"""Branch coverage for docx_report.generate_ombudsman_docx dispatcher.

Targets the full report-generation pipeline that single-section unit tests
cannot reach: the section dispatch loop, per-section error isolation,
financial-total helpers, RenderContext wiring, the GUI wrapper, and the
edge guards (empty records / missing engine / absent columns).
"""

from __future__ import annotations

from typing import Any

import pytest
from docx import Document

from edf_bill_fetcher.io.reporters.docx_report import (
    generate_docx_from_gui,
    generate_ombudsman_docx,
)


def _make_records(n: int, *, seed: int = 0) -> list[dict]:
    rows = []
    for i in range(n):
        month = ((i + seed) % 12) + 1
        year = 2018 + ((i + seed) // 12)
        day = 1 + ((i + seed) % 27)
        rows.append(
            {
                "Date": f"{day:02d}/{month:02d}/{year}",
                "Source": "Local PDF Folder",
                "Entry Type": "New Bill" if i % 5 else "Payment",
                "Amount (£)": 50.0 + (i % 7) * 1.5,
                "Period Charge (£)": 80.0 if i % 5 else 0.0,
                "Period From": f"01/{month:02d}/{year}",
                "Period To": f"28/{month:02d}/{year}",
                "Invoice #": f"INV-{i:05d}",
                "Reading": "Actual",
                "Units (kWh)": "100",
                "Standing Chg (p/day)": "60.1",
                "Tariff": "Flexible Octopus",
                "Attachment Name": f"bill_{i:05d}.pdf",
                "Details": "Auto-debit via DD reference 1234567",
                "Logic Used": "Period Charge Match",
            }
        )
    return rows


class _StubEngine:
    def __init__(self) -> None:
        self.pdf_count = 3
        self.email_count = 2
        self.filtered_records: list[dict] = [
            {
                "Source": "PST",
                "Date": "01/01/2020",
                "Amount (£)": 12.0,
                "Details": "below threshold",
            }
        ]


def test_generate_full_report_all_sections(tmp_path: Any) -> None:
    records = _make_records(24)
    out = str(tmp_path / "full.docx")
    config = {"report_account_ref": "A-12345678", "report_sections": []}
    result = generate_ombudsman_docx(records, out, config, _StubEngine())
    assert result == out
    doc = Document(out)
    assert len(doc.paragraphs) > 10


def test_generate_subset_of_sections(tmp_path: Any) -> None:
    records = _make_records(12)
    out = str(tmp_path / "subset.docx")
    config = {"acc_num": "A-87654321", "report_sections": ["exec_summary", "timeline"]}
    result = generate_ombudsman_docx(records, out, config, None)
    assert result == out


def test_generate_with_filtered_records(tmp_path: Any) -> None:
    records = _make_records(6)
    out = str(tmp_path / "filtered.docx")
    config = {"report_account_ref": "A-11111111", "report_sections": []}
    result = generate_ombudsman_docx(records, out, config, None, filtered=[{"x": 1}])
    assert result == out


def test_generate_empty_records_raises(tmp_path: Any) -> None:
    with pytest.raises(ValueError, match="No records"):
        generate_ombudsman_docx([], str(tmp_path / "x.docx"), {}, None)


def test_generate_section_failure_is_isolated(
    tmp_path: Any, monkeypatch: pytest.MonkeyPatch
) -> None:
    """A broken section builder degrades to a visible paragraph, not an abort."""
    import edf_bill_fetcher.io.reporters.docx_report as docx_mod

    def _boom(**kwargs: Any) -> None:
        raise RuntimeError("section exploded")

    monkeypatch.setattr(docx_mod, "create_executive_summary", _boom)
    records = _make_records(6)
    out = str(tmp_path / "isolated.docx")
    config = {"report_account_ref": "A-1", "report_sections": []}
    result = generate_ombudsman_docx(records, out, config, None)
    assert result == out
    doc = Document(out)
    texts = [p.text for p in doc.paragraphs]
    assert any("Executive Summary failed" in t for t in texts)


def test_generate_docx_from_gui_success(tmp_path: Any) -> None:
    records = _make_records(6)
    out = str(tmp_path / "gui.docx")
    ok, msg = generate_docx_from_gui(records, out, {"report_account_ref": "A-2"}, None)
    assert ok is True
    assert "generated" in msg.lower()


def test_generate_docx_from_gui_failure() -> None:
    ok, msg = generate_docx_from_gui([], "/nonexistent/x.docx", {}, None)
    assert ok is False
    assert "Failed" in msg
