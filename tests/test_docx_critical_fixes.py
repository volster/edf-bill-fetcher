"""Regression tests pinning the three CRITICAL bugs flagged in the dev-branch
code review (2026-07). Each of these shipped silently to end users because
there was no test guarding the visual/contract surface:

- DocumentCoverPageLabelVisibility — the cover-page label table is rendered
  with bold NAVY/bold text in row 0; the pre-fix ``_format_table`` call
  silently overwrote that with white-bold on light-blue fill (invisible).
  Fixed by adding ``header_row=False`` to that call site.

- DocumentGlossaryHeaderRowPopulated — the DOCX glossary pre-fix allocated
  rows for header + terms but never wrote the header cells (it started the
  loop at index 1), then ``_format_table`` painted the empty row 0 with the
  header shading.  Pre-fix output: blank header row.  Fixed by writing
  ``["Term", "Definition"]`` into row 0 before the iteration.

- DocumentOFGEMAutoCarryForward — the DOCX OFGEM section pre-fix only
  inspected quarters present in the ``_load_ofgem_caps`` table; quarters
  beyond it (e.g. 2026-Q4 when the cap table only lists up to 2026-Q3) were
  silently marked ``CAP DATA UNAVAILABLE``.  The PDF surfaces the
  ``_LATEST_KNOWN`` carry-forward sentinel and routes to a
  ``COMPLIANT (CARRIED)`` overall verdict.  Fixed by porting the PDF
  carry-forward branch verbatim, including reading the sentinel and adding
  the ``elif carried_count > 0`` summary row.

Kept deliberately stand-alone (not folded into ``test_report.py``) to keep
the diff straightforward and to make a future regression back into a single
bug class easy to triage.
"""

from __future__ import annotations

import sys
from typing import Any

import pandas as pd
import pytest
from docx import Document
from docx.document import Document as DocumentType

from edf_bill_fetcher.io.reporters.docx_report import (
    NAVY,
    WHITE,
    _format_table,
    _get_or_create_styles,
    create_appendix_glossary,
    create_cover_page,
    create_ofgem_comparison,
)

# ---------------------------------------------------------------------------
# Fixtures
# ---------------------------------------------------------------------------


@pytest.fixture
def doc_styles() -> tuple[DocumentType, Any]:
    d = Document()
    return d, _get_or_create_styles(d)


def _runs_text(runs: Any) -> str:
    return "".join(r.text for r in runs)


def _row_cell_text(cell: Any) -> str:
    return _runs_text(p.runs[0] for p in cell.paragraphs if p.runs)


# ---------------------------------------------------------------------------
# Cover-page label visibility
# ---------------------------------------------------------------------------


class TestCoverPageLabelVisibility:
    """Pins the CRITICAL DOCX bug: cover-page labels were rendered
    white-bold on a light-blue fill and effectively invisible.
    """

    def test_account_reference_label_renders_in_navy_not_white(
        self, doc_styles: tuple[DocumentType, Any]
    ) -> None:
        doc, styles = doc_styles
        create_cover_page(
            doc,
            styles,
            acc_ref="A-12345678",
            period_start="01 Jan 2024",
            period_end="31 Jan 2024",
            report_date="15 Jan 2024",
        )
        # The cover-page table is the second-table-on-the-page (after the
        # document does not add anything before it).  Pull its row 0 +
        # take the first cell — that is the "Account Reference" label.
        # Allow the table to be the only OR the last-added table; the
        # _format_table() on behalf of the cover now leaves row 0 alone.
        cover_table = doc.tables[-1]
        label_cell = cover_table.rows[0].cells[0]
        # The label must read "Account Reference" in NAVY, not WHITE.
        # Pre-fix the run was forced to WHITE by _format_table.
        assert _row_cell_text(label_cell).strip() == "Account Reference"
        assert label_cell.paragraphs[0].runs[0].font.color.rgb == NAVY
        assert label_cell.paragraphs[0].runs[0].font.color.rgb != WHITE
        # And the data cell (column 1) must NOT be repainted as a header.
        value_cell = cover_table.rows[0].cells[1]
        assert value_cell.paragraphs[0].runs[0].font.color.rgb != WHITE

    def test_format_table_header_row_false_opt_out(self) -> None:
        """``_format_table`` must accept ``header_row=False`` and skip the
        header-row repaint.  Mirrors the call-site fix for the cover page.
        """
        d = Document()
        t = d.add_table(rows=2, cols=2)
        t.rows[0].cells[0].text = "LABEL"
        t.rows[0].cells[1].text = "value"
        t.rows[1].cells[0].text = "other label"
        t.rows[1].cells[1].text = "other value"
        _format_table(t, header_color="#EBF3FA", font_size=10, header_row=False)

        # Row 0 cell must NOT have been repainted white; the cover-page
        # label styling should remain intact.
        assert t.rows[0].cells[0].paragraphs[0].runs[0].font.color.rgb != WHITE
        # The data-cell styling path still applies (DARK_GREY) — at minimum,
        # tag values are not WHITE either.
        assert t.rows[0].cells[1].paragraphs[0].runs[0].font.color.rgb != WHITE

    def test_format_table_header_row_true_default_keeps_legacy(self) -> None:
        """When ``header_row`` is left default True, the function must
        continue to mark row 0 cells white-bold as before so consumers
        that explicitly want a header row get one.
        """
        d = Document()
        t = d.add_table(rows=1, cols=1)
        t.rows[0].cells[0].text = "Header"
        _format_table(t, header_color="#10367A", font_size=8)  # default header_row=True
        assert t.rows[0].cells[0].paragraphs[0].runs[0].font.color.rgb == WHITE


# ---------------------------------------------------------------------------
# Glossary header row
# ---------------------------------------------------------------------------


class TestGlossaryHeaderRowPopulated:
    """Pins the CRITICAL DOCX bug: glossary row 0 was never populated, so
    the rendered PDF had a blank header row with header shading.
    """

    def test_glossary_row_0_has_term_and_definition(self) -> None:
        d = Document()
        styles = _get_or_create_styles(d)
        create_appendix_glossary(d, styles)
        gtable = d.tables[-1]
        # Row 0 must now read the actual two-column header, not be blank.
        assert _row_cell_text(gtable.rows[0].cells[0]).strip() == "Term"
        assert _row_cell_text(gtable.rows[0].cells[1]).strip() == "Definition"

    def test_glossary_body_rows_present_below_header(self) -> None:
        """Pre-fix the loop started at index 1 so row 0 stayed empty.  Post-fix
        the loop still iterates the terms starting at index 1, leaving the
        terms populated below the now-correct header.
        """
        d = Document()
        styles = _get_or_create_styles(d)
        create_appendix_glossary(d, styles)
        gtable = d.tables[-1]
        # First body row must be one of the glossary term keys (the dict
        # ordering dictates the first entry; we don't pin the order, just
        # that the row is non-empty and not a duplicated header).
        first_body = _row_cell_text(gtable.rows[1].cells[0]).strip()
        assert first_body != ""
        assert first_body != "Term"
        assert _row_cell_text(gtable.rows[1].cells[1]).strip() != ""


# ---------------------------------------------------------------------------
# OFGEM carry-forward
# ---------------------------------------------------------------------------


def _install_minimal_caps(monkeypatch: pytest.MonkeyPatch, *, with_latest: bool) -> None:
    """Install a deterministic one-quarter OFGEM cap table for the test
    plus, optionally, a carry-forward cap as the tuple's second element.

    ``_load_ofgem_caps`` is bound into ``docx_report``'s module
    namespace at import time (via ``from edf_bill_fetcher.io.reporters.pdf_report import _load_ofgem_caps``
    on line 34 of ``docx_report.py``) and the call site resolves it
    locally — patching
    ``sys.modules["edf_bill_fetcher.io.reporters.pdf_report"]._load_ofgem_caps`` alone is therefore
    insufficient; we must patch the symbol in the ``docx_report``
    namespace too, which is what this helper does.
    """
    minimal_caps: dict[str, dict[str, float]] = {
        "2026-Q3": {"unit_rate": 25.0},
    }
    latest = {"unit_rate": 25.0} if with_latest else None
    monkeypatch.setattr(
        sys.modules["edf_bill_fetcher.io.reporters.docx_report"],
        "_load_ofgem_caps",
        lambda auto_carry=False: (minimal_caps, latest),
    )


def _make_ofgem_records(quarter_year: int, quarter_num: int) -> pd.DataFrame:
    """Synthesise a one-row DataFrame matching ``_period_to_ofgem_quarter``'s
    expectations: an actual Date in the quarter, Period Charge + Units
    populated, so the builder processes the row through to the quarterly
    comparison step.
    """
    return pd.DataFrame(
        [
            {
                "Date": f"15 {['Jan', 'Apr', 'Jul', 'Oct'][quarter_num - 1]} {quarter_year}",
                "Source": "Local PDF Folder",
                "Entry Type": "New Bill",
                "Amount (£)": 200.0,
                "Period Charge (£)": 200.0,
                "Units (kWh)": 800.0,  # 200.0 / 800 * 100 = 25.0 p/kWh
                "Period From": f"01 {['Jan', 'Apr', 'Jul', 'Oct'][quarter_num - 1]} {quarter_year}",
                "Period To": f"30 {['Jan', 'Apr', 'Jul', 'Oct'][quarter_num - 1]} {quarter_year}",
                "Invoice #": "TEST-001",
                "Reading": "Actual",
            }
        ]
    )


class TestOFGEMAutoCarryForward:
    """Pins the CRITICAL DOCX bug: quarters beyond the published OFGEM cap
    table were silently marked ``CAP DATA UNAVAILABLE`` and produced an
    ``INCOMPLETE`` overall verdict in the DOCX, even though the PDF surfaces
    the carry-forward cap path with a ``COMPLIANT (CARRIED)``
    verdict.
    """

    def test_quarter_beyond_table_routes_to_carry_forward(
        self, monkeypatch: pytest.MonkeyPatch
    ) -> None:
        """With the carry-forward cap installed, a quarter BEYOND the
        hard-coded table (we manufacture one) must be presented as a
        carried-forward row (``EXCEEDS CAP (CAP CARRIED FORWARD)`` / ``AT
        CAP (CAP CARRIED FORWARD)`` / ``BELOW CAP (CAP CARRIED FORWARD)``)
        and the summary row must read ``COMPLIANT (CARRIED)``.
        """
        _install_minimal_caps(monkeypatch, with_latest=True)
        d = Document()
        styles = _get_or_create_styles(d)
        df = _make_ofgem_records(quarter_year=2026, quarter_num=4)  # Q4 not in cap table
        create_ofgem_comparison(d, styles, df, config={})

        # Find the comparison table: it's the one whose row 0 is the
        # ``Period / Bill Unit Rate / OFGEM Cap / Difference / Status``
        # header.  There will also be the OFGEM caps reference table later.
        comp_table = None
        for t in d.tables:
            if len(t.rows) >= 2 and len(t.columns) == 5:
                hdr = [t.rows[0].cells[j].text.strip() for j in range(5)]
                if hdr == [
                    "Period",
                    "Bill Unit Rate (p/kWh)",
                    "OFGEM Cap (p/kWh)",
                    "Difference",
                    "Status",
                ]:
                    comp_table = t
                    break
        assert comp_table is not None, "comparison table not found"

        # The body row for the carried-forward quarter.
        body = comp_table.rows[1]
        assert body.cells[0].text.strip() == "2026-Q4"
        # Cap Rate / Difference should be numeric (carry-forward values),
        # not the MISSING sentinel "—".
        assert body.cells[2].text.strip() != "—"
        # Status should reference the carried-forward marker.
        assert "CAP CARRIED FORWARD" in body.cells[4].text

        # The summary row at index len(cap_rows) + 1 should read
        # ``COMPLIANT (CARRIED)`` (the carried-forward verdict), not
        # ``INCOMPLETE`` (which the pre-fix DOCX emitted when ANY quarter
        # was unavailable).
        summary = comp_table.rows[2]
        assert summary.cells[0].text.strip() == "OVERALL"
        assert "COMPLIANT (CARRIED)" in summary.cells[4].text

    def test_quarter_beyond_table_without_latest_known_still_incomplete(
        self, monkeypatch: pytest.MonkeyPatch
    ) -> None:
        """When the cap table is installed WITHOUT a carry-forward cap (i.e.
        a future-quarter fallback would fail) the DOCX must still fall
        back to ``INCOMPLETE`` — the carry-forward pin above does not
        silently invent a cap.
        """
        _install_minimal_caps(monkeypatch, with_latest=False)
        d = Document()
        styles = _get_or_create_styles(d)
        df = _make_ofgem_records(quarter_year=2026, quarter_num=4)
        create_ofgem_comparison(d, styles, df, config={})

        # Find the comparison table again.
        comp_table = None
        for t in d.tables:
            if len(t.rows) >= 2 and len(t.columns) == 5:
                hdr = [t.rows[0].cells[j].text.strip() for j in range(5)]
                if hdr == [
                    "Period",
                    "Bill Unit Rate (p/kWh)",
                    "OFGEM Cap (p/kWh)",
                    "Difference",
                    "Status",
                ]:
                    comp_table = t
                    break
        assert comp_table is not None

        # Body row keeps the row but flags UNAVAILABLE (we want the
        # reviewer to see the quarter exists in the data).
        body = comp_table.rows[1]
        assert body.cells[2].text.strip() == "—"
        assert "UNAVAILABLE" in body.cells[4].text

        # Overall verdict is INCOMPLETE.
        summary = comp_table.rows[2]
        assert summary.cells[0].text.strip() == "OVERALL"
        assert "INCOMPLETE" in summary.cells[4].text

    def test_quarter_with_nan_unit_rate_shows_na_row(
        self, monkeypatch: pytest.MonkeyPatch
    ) -> None:
        """A quarter whose unit rate is NaN (0/0) must render an "N/A" row
        rather than silently disappearing from the comparison table (L-12),
        so the reader knows the quarter existed in the data.
        """
        _install_minimal_caps(monkeypatch, with_latest=True)
        d = Document()
        styles = _get_or_create_styles(d)
        df = pd.DataFrame(
            [
                {
                    "Date": "15 Oct 2026",  # 2026-Q4, beyond the cap table
                    "Source": "Local PDF Folder",
                    "Entry Type": "New Bill",
                    "Amount (£)": 200.0,
                    "Period Charge (£)": 0.0,
                    "Units (kWh)": 0.0,  # 0/0 → NaN unit rate
                    "Period From": "01 Oct 2026",
                    "Period To": "30 Oct 2026",
                    "Invoice #": "TEST-002",
                    "Reading": "Actual",
                }
            ]
        )
        create_ofgem_comparison(d, styles, df, config={})

        comp_table = None
        for t in d.tables:
            if len(t.rows) >= 2 and len(t.columns) == 5:
                hdr = [t.rows[0].cells[j].text.strip() for j in range(5)]
                if hdr == [
                    "Period",
                    "Bill Unit Rate (p/kWh)",
                    "OFGEM Cap (p/kWh)",
                    "Difference",
                    "Status",
                ]:
                    comp_table = t
                    break
        assert comp_table is not None, "comparison table not found"

        # The quarter must still appear, with the bill rate as "N/A".
        body = comp_table.rows[1]
        assert body.cells[0].text.strip() == "2026-Q4"
        assert body.cells[1].text.strip() == "N/A"
        # The carried-forward cap is still shown; difference is unknown.
        assert body.cells[2].text.strip() == "25.00"
        assert body.cells[3].text.strip() == "N/A"
        assert body.cells[4].text.strip() == "N/A"

    def test_signature_accepts_config(self) -> None:
        """Regression-pin: ``create_ofgem_comparison`` must accept a
        ``config: dict | None`` argument (matches the PDF signature —
        cross-format symmetry).  Pre-fix the DOCX signature was
        ``(doc, styles, df, ctx)``; post-fix it is
        ``(doc, styles, df, config=None, ctx=None)``.
        """
        import inspect

        sig = inspect.signature(create_ofgem_comparison)
        params = list(sig.parameters.keys())
        assert "config" in params, (
            f"create_ofgem_comparison must accept 'config' for PDF-parity; got parameters {params}"
        )
        # config must default to None so existing call sites stay valid.
        assert sig.parameters["config"].default is None


# ---------------------------------------------------------------------------
# Sanity guard: test_dispatch_parity-style structural walk — bonus
# ---------------------------------------------------------------------------


def test_format_table_is_idempotent_on_label_value_table() -> None:
    """Sanity check: calling ``_format_table`` twice on a label/value table
    with ``header_row=False`` does not corrupt the runs.  Avoids accidental
    regressions in the helper.
    """
    d = Document()
    t = d.add_table(rows=2, cols=2)
    t.rows[0].cells[0].text = "Label"
    t.rows[0].cells[1].text = "value"
    _format_table(t, header_color="#EBF3FA", font_size=10, header_row=False)
    # Snapshot — the helper should not corrupt by repainting.
    color1 = t.rows[0].cells[0].paragraphs[0].runs[0].font.color.rgb
    _format_table(t, header_color="#EBF3FA", font_size=10, header_row=False)
    color2 = t.rows[0].cells[0].paragraphs[0].runs[0].font.color.rgb
    assert color1 == color2


def test_format_table_default_header_row_true_is_preserved() -> None:
    """The default ``header_row=True`` is what every existing call site
    relied on before this fix landed; make sure the default hasn't
    silently flipped to False, which would break the OFGEM / evidence /
    methodology tables throughout the report.
    """
    import inspect

    sig = inspect.signature(_format_table)
    assert sig.parameters["header_row"].default is True
