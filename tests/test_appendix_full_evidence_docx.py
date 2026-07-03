from __future__ import annotations

from typing import Any

import pandas as pd
import pytest
from docx import Document
from docx.document import Document as DocumentType

from edf_report_docx import _get_or_create_styles, create_appendix_full_evidence


@pytest.fixture
def states() -> tuple[DocumentType, Any]:
    doc = Document()
    styles = _get_or_create_styles(doc)
    return doc, styles


def _make_records(n: int) -> list[dict]:
    rows = []
    for i in range(n):
        month = (i % 12) + 1
        year = 2018 + (i // 12)
        rows.append(
            {
                "Date": f"{1 + (i % 27):02d}/{month:02d}/{year}",
                "Source": "Local PDF Folder",
                "Entry Type": "New Bill" if i % 5 else "Payment",
                "Amount (£)": 50.0 + (i % 7) * 1.5,
                "Period Charge (£)": 80.0 if i % 5 else 0.0,
                "Period From": f"01/{month:02d}/{year}",
                "Period To": f"28/{month:02d}/{year}",
                "Invoice #": f"INV-{i:05d}",
                "Reading": "Actual",
                "Units (kWh)": "100",
                "Standing Chg (p/day)": "60.1",
                "Attachment Name": f"bill_{i:05d}.pdf",
                "Details": "Auto-debit via DD reference 1234567",
            }
        )
    return rows


class TestAppendixCap:
    def test_small_dataset_unchanged(self, states: tuple[DocumentType, Any]) -> None:
        doc, styles = states
        records = _make_records(20)
        df = pd.DataFrame(records)
        result = create_appendix_full_evidence(doc, styles, df)
        assert result["rendered_rows"] == 20
        assert result["total_rows"] == 20
        assert result["truncated"] is False

    def test_at_cap_does_not_truncate(self, states: tuple[DocumentType, Any]) -> None:
        doc, styles = states
        records = _make_records(150)
        df = pd.DataFrame(records)
        result = create_appendix_full_evidence(doc, styles, df)
        # At the boundary, the cap is inclusive: 150 rows render
        # all 150.
        assert result["rendered_rows"] == 150
        assert result["total_rows"] == 150
        assert result["truncated"] is False

    def test_over_cap_truncates(self, states: tuple[DocumentType, Any]) -> None:
        doc, styles = states
        records = _make_records(500)
        df = pd.DataFrame(records)
        result = create_appendix_full_evidence(doc, styles, df)
        assert result["rendered_rows"] == 150
        assert result["total_rows"] == 500
        assert result["truncated"] is True

    def test_truncation_note_is_in_document(self, states: tuple[DocumentType, Any]) -> None:
        doc, styles = states
        records = _make_records(500)
        df = pd.DataFrame(records)
        create_appendix_full_evidence(doc, styles, df)
        body = "\n".join(p.text for p in doc.paragraphs)
        # The truncation explanation's wording matches the user-
        # approved text in the Phase 2.3 spec.
        assert "Please refer to the accompanying" in body
        assert "Excel workbook" in body

    def test_chronological_first_rows_kept(self, states: tuple[DocumentType, Any]) -> None:
        # Records come in date-ascending order, so the oldest
        # row appears first after the truncation slice.
        doc, styles = states
        records = _make_records(300)
        df = pd.DataFrame(records)
        result = create_appendix_full_evidence(doc, styles, df)
        assert result["truncated"] is True
        assert result["rendered_rows"] == 150
        # The first rendered data row should be a "01/..." date
        # suffix because the synthesised fixture starts with
        # day=1 of each month.
        first_table = doc.tables[-1]
        first_data_row = first_table.rows[1]
        first_cell = first_data_row.cells[0].text
        assert first_cell.startswith("01/")
