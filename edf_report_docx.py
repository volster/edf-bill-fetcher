"""
DOCX Report Generator for EDF Energy Ombudsman Submissions.

Generates a professional Word document report optimized for Energy Ombudsman review.
Mirrors the PDF report structure but uses python-docx for output.
"""

from __future__ import annotations

from datetime import datetime
from typing import Any

import numpy as np
import pandas as pd
from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import parse_xml
from docx.oxml.ns import qn
from docx.shared import (
    Cm,
    Pt,
    RGBColor,
)

# Import from main module
from edf_collector import parse_to_sort_date
from edf_report import (
    REPORT_SECTIONS,
    RenderContext,
    _compute_mean_daily,
    _get_package_version,
    _load_ofgem_caps,
)

# =============================================================================
# CONSTANTS
# =============================================================================

NAVY = RGBColor(0x10, 0x36, 0x7A)
DARK_BLUE = RGBColor(0x1B, 0x4F, 0x9E)
MEDIUM_BLUE = RGBColor(0x2E, 0x75, 0xB6)
LIGHT_BLUE = RGBColor(0xD6, 0xE4, 0xF0)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
BLACK = RGBColor(0x00, 0x00, 0x00)
DARK_GREY = RGBColor(0x33, 0x33, 0x33)
MEDIUM_GREY = RGBColor(0x66, 0x66, 0x66)
LIGHT_GREY = RGBColor(0xF2, 0xF2, 0xF2)
RED = RGBColor(0xC0, 0x00, 0x00)
AMBER = RGBColor(0xED, 0x7D, 0x31)
GREEN = RGBColor(0x54, 0x82, 0x35)
ORANGE = RGBColor(0xFE, 0x57, 0x16)

MARGIN_CM = 2.5


def fmt_date(val: Any) -> str:
    """Format date for display."""
    if val is None or pd.isna(val):
        return "Unknown"
    if hasattr(val, "strftime"):
        return val.strftime("%d %B %Y")  # type: ignore[no-any-return]
    return str(val)


def fmt_number(val: Any, decimals: int = 2) -> str:
    """Format number with commas."""
    if val is None or pd.isna(val):
        return "N/A"
    return f"{float(val):,.{decimals}f}"


def fmt_money(val: Any) -> str:
    """Format money with £ sign.

    Signed-zero guard: a value like ``-0.001`` rounds in f-strings to
    ``£-0.00``, which looks wrong on a report.  Any value whose absolute
    value rounds to zero at 2 decimal places is coerced to plain ``0.0``
    before formatting, matching the PDF generator (edf_report.py:254).
    """
    if val is None or pd.isna(val):
        return "N/A"
    try:
        f = float(val)
        if abs(f) < 0.005:  # rounds to 0.00 at 2-dp display
            f = 0.0
        return f"£{f:,.2f}"
    except (ValueError, TypeError):
        return str(val)


# =============================================================================
# STYLE HELPERS
# =============================================================================


def _add_heading_style(
    doc: Any,
    level: int,
    name: str,
    font_size: int,
    color: RGBColor = NAVY,
    bold: bool = True,
    space_before: int = 12,
    space_after: int = 6,
) -> Any:
    """Add a heading style to the document."""
    style = doc.styles.add_style(name, 1)  # WD_STYLE_TYPE.PARAGRAPH
    style.font.size = Pt(font_size)
    style.font.color.rgb = color
    style.font.bold = bold
    style.paragraph_format.space_before = Pt(space_before)
    style.paragraph_format.space_after = Pt(space_after)
    style.paragraph_format.keep_with_next = True
    return style


def _get_or_create_styles(doc: Any) -> Any:
    """Create custom styles for the document."""
    # Title
    if "CoverTitle" not in doc.styles:
        s = doc.styles.add_style("CoverTitle", 1)
        s.font.size = Pt(28)
        s.font.color.rgb = NAVY
        s.font.bold = True
        s.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        s.paragraph_format.space_after = Pt(12)

    # Subtitle
    if "CoverSubtitle" not in doc.styles:
        s = doc.styles.add_style("CoverSubtitle", 1)
        s.font.size = Pt(14)
        s.font.color.rgb = MEDIUM_BLUE
        s.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        s.paragraph_format.space_after = Pt(6)

    # Section header
    if "SectionHeader" not in doc.styles:
        s = doc.styles.add_style("SectionHeader", 1)
        s.font.size = Pt(16)
        s.font.color.rgb = NAVY
        s.font.bold = True
        s.paragraph_format.space_before = Pt(18)
        s.paragraph_format.space_after = Pt(10)

    # Sub-section header
    if "SubSectionHeader" not in doc.styles:
        s = doc.styles.add_style("SubSectionHeader", 1)
        s.font.size = Pt(13)
        s.font.color.rgb = DARK_BLUE
        s.font.bold = True
        s.paragraph_format.space_before = Pt(12)
        s.paragraph_format.space_after = Pt(6)

    # Body text
    if "BodyText" not in doc.styles:
        s = doc.styles.add_style("BodyText", 1)
        s.font.size = Pt(11)
        s.font.color.rgb = DARK_GREY
        s.paragraph_format.space_after = Pt(6)
        s.paragraph_format.line_spacing = 1.15

    # Caption
    if "Caption" not in doc.styles:
        s = doc.styles.add_style("Caption", 1)
        s.font.size = Pt(9)
        s.font.color.rgb = MEDIUM_GREY
        s.font.italic = True
        s.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        s.paragraph_format.space_before = Pt(4)
        s.paragraph_format.space_after = Pt(8)

    # Footer
    if "FooterText" not in doc.styles:
        s = doc.styles.add_style("FooterText", 1)
        s.font.size = Pt(8)
        s.font.color.rgb = MEDIUM_GREY
        s.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

    return {
        "CoverTitle": "CoverTitle",
        "CoverSubtitle": "CoverSubtitle",
        "SectionHeader": "SectionHeader",
        "SubSectionHeader": "SubSectionHeader",
        "BodyText": "BodyText",
        "Caption": "Caption",
        "FooterText": "FooterText",
    }


def _add_footer(doc: Any, text: str = "EDF Energy Billing Evidence Report — Confidential") -> None:
    """Add footer to all sections."""
    for section in doc.sections:
        footer = section.footer
        footer.is_linked_to_previous = False
        p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        p.style = doc.styles["FooterText"]
        p.text = text


def _set_cell_shading(cell, color):
    """Set background color for table cell. Accepts hex string (e.g., '#FF0000') or RGBColor object."""
    # Convert RGBColor to hex string if needed
    if hasattr(color, "rgb"):
        # RGBColor object
        hex_str = f"{color.rgb:06X}"
    elif isinstance(color, str):
        # Hex string
        hex_str = color.lstrip("#")
    else:
        # Fallback
        hex_str = str(color).lstrip("#")

    shading_elm = cell._element.get_or_add_tcPr()
    shading = shading_elm.makeelement(
        qn("w:shd"),
        {
            qn("w:fill"): hex_str,
            qn("w:val"): "clear",
        },
    )
    shading_elm.append(shading)


def _format_table(table, header_color="#10367A", font_size=8):
    """Apply formatting to a docx table."""
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header row
    if table.rows:
        for cell in table.rows[0].cells:
            _set_cell_shading(cell, header_color)
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in paragraph.runs:
                    run.font.color.rgb = WHITE
                    run.font.bold = True
                    run.font.size = Pt(font_size)

    # Data rows
    for i, row in enumerate(table.rows[1:], 1):
        for cell in row.cells:
            if i % 2 == 0:
                _set_cell_shading(cell, LIGHT_GREY)
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(font_size)
                    run.font.color.rgb = DARK_GREY


# =============================================================================
# SECTION CREATORS
# =============================================================================


def create_cover_page(
    doc: Any,
    styles: Any,
    acc_ref: str,
    period_start: str,
    period_end: str,
    report_date: str,
) -> None:
    """Create cover page elements."""
    doc.add_paragraph("EDF ENERGY BILLING", style=styles["CoverTitle"])
    doc.add_paragraph("EVIDENCE REPORT", style=styles["CoverTitle"])
    doc.add_paragraph("")  # spacer

    p = doc.add_paragraph(style=styles["CoverSubtitle"])
    p.add_run("Prepared for Energy Ombudsman Review").bold = True

    doc.add_paragraph("")

    # Info table
    table = doc.add_table(rows=5, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    info = [
        ("Account Reference", acc_ref),
        ("Report Period", f"{period_start} to {period_end}"),
        ("Report Date", report_date),
        ("Generated By", f"EDF Bill Fetcher v{_get_package_version()}"),
        ("Classification", "CONFIDENTIAL — For Ombudsman Use Only"),
    ]

    for i, (label, value) in enumerate(info):
        row = table.rows[i]
        row.cells[0].text = label
        row.cells[1].text = value
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(11)
                    if cell == row.cells[0]:
                        run.font.bold = True
                        run.font.color.rgb = NAVY
                    else:
                        run.font.color.rgb = DARK_GREY

    _format_table(table, header_color="#EBF3FA", font_size=11)

    doc.add_paragraph("")
    doc.add_page_break()


def create_table_of_contents(doc: Any, styles: Any, ctx: RenderContext | None = None) -> None:
    """Create a single-column TOC driven by the registry.

    Same numbering rules as ``edf_report.create_table_of_contents``: main
    sections numbered 1, 2, 3...; appendices lettered A, B, C...; both
    derived from REPORT_SECTIONS plus the user's selected_sections.
    """
    doc.add_paragraph("TABLE OF CONTENTS", style=styles["SectionHeader"])

    if ctx is None:
        ctx = RenderContext()
    sections = [(spec.label, spec.section.title) for spec in ctx.sections_in_order]

    if not sections:
        p = doc.add_paragraph(style=styles["BodyText"])
        p.add_run("No report sections selected.").italic = True
        doc.add_page_break()
        return

    for label, title in sections:
        p = doc.add_paragraph(style=styles["BodyText"])
        run = p.add_run(f"{label} {title}")
        run.font.size = Pt(11)
        run.font.bold = True

    doc.add_page_break()


def _add_banner_heading(doc: Any, text: str) -> None:
    doc.add_paragraph(text, style="Heading 1")


def create_executive_summary(
    doc: Any,
    styles: Any,
    df: pd.DataFrame,
    config: dict,
    acc_ref: str,
    flag_counts: dict,
    n_records: int,
    charges: float,
    payments: float,
    period_start: str,
    period_end: str,
    ctx: RenderContext | None = None,
) -> None:
    """Create executive summary section."""
    if ctx is None:
        ctx = RenderContext()
    doc.add_paragraph(ctx.heading("exec_summary"), style=styles["SectionHeader"])

    doc.add_paragraph(
        f"This report presents a comprehensive analysis of EDF Energy billing records "
        f"for account {acc_ref}, covering the period {period_start} to {period_end}. "
        f"The analysis is based on {n_records} billing records extracted from "
        f"PDF invoices, HTM account history exports, and PST/OST email archives.",
        style=styles["BodyText"],
    )

    # Financial snapshot table
    table = doc.add_table(rows=5, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    exec_data = [
        ("Total Records Analyzed", str(n_records)),
        ("Period Covered", f"{period_start} – {period_end}"),
        ("Total Charges (Debits)", fmt_money(charges)),
        ("Total Payments (Credits)", fmt_money(payments)),
        ("Net Position", fmt_money(charges - payments)),
    ]

    for i, (label, value) in enumerate(exec_data):
        row = table.rows[i]
        row.cells[0].text = label
        row.cells[1].text = value
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(11)
                    if cell == row.cells[0]:
                        run.font.bold = True
                        run.font.color.rgb = NAVY
                    else:
                        run.font.color.rgb = DARK_GREY

    _format_table(table, header_color="#EBF3FA", font_size=11)

    doc.add_paragraph("")

    # Key metrics
    doc.add_paragraph("Key Metrics:", style=styles["SubSectionHeader"])

    if "Amount (£)" in df.columns:
        amounts = pd.to_numeric(df["Amount (£)"], errors="coerce").dropna()
        if not amounts.empty:
            doc.add_paragraph(
                f"• Average bill amount: {fmt_money(amounts.mean())}",
                style=styles["BodyText"],
            )
            doc.add_paragraph(
                f"• Median bill amount: {fmt_money(amounts.median())}",
                style=styles["BodyText"],
            )
            doc.add_paragraph(
                f"• Highest single charge: {fmt_money(amounts.max())}",
                style=styles["BodyText"],
            )
            doc.add_paragraph(
                f"• Lowest single charge: {fmt_money(amounts.min())}",
                style=styles["BodyText"],
            )

    doc.add_paragraph("")


def create_key_findings_table(
    doc: Any, styles: Any, flags: list, ctx: RenderContext | None = None
) -> None:
    """Create key findings summary table from flags."""
    if ctx is None:
        ctx = RenderContext()
    doc.add_paragraph(ctx.heading("key_findings"), style=styles["SectionHeader"])

    if not flags:
        doc.add_paragraph(
            "No automated flags were raised. Manual review of the evidence index and "
            "detailed findings sections is recommended.",
            style=styles["BodyText"],
        )
    else:
        table = doc.add_table(rows=len(flags) + 1, cols=4)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        # Header
        headers = ["Severity", "Category", "Description", "Records Affected"]
        for i, h in enumerate(headers):
            table.rows[0].cells[i].text = h

        for i, flag in enumerate(flags, 1):
            row = table.rows[i]
            # Flag tuple: (type, date, amount, detail, severity)
            row.cells[0].text = flag[4] if len(flag) > 4 else ""  # Severity
            row.cells[1].text = flag[0] if len(flag) > 0 else ""  # Category (type)
            row.cells[2].text = flag[3] if len(flag) > 3 else ""  # Description (detail)
            row.cells[3].text = (
                fmt_date(flag[1]) if len(flag) > 1 else ""
            )  # Records Affected (date)

        _format_table(table)

    doc.add_page_break()


def create_evidence_index(
    doc: Any,
    styles: Any,
    df: pd.DataFrame,
    engine: Any,
    ctx: RenderContext | None = None,
) -> None:
    """Create evidence index with source cross-references."""
    if ctx is None:
        ctx = RenderContext()
    doc.add_paragraph(ctx.heading("evidence_index"), style=styles["SectionHeader"])

    doc.add_paragraph(
        "The following table summarizes all source documents processed during extraction.",
        style=styles["BodyText"],
    )

    # Source breakdown
    source_counts = df["Source"].value_counts()

    table = doc.add_table(rows=len(source_counts) + 2, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    headers = ["Source Type", "Records", "Percentage"]
    for i, h in enumerate(headers):
        table.rows[0].cells[i].text = h

    for i, (src, cnt) in enumerate(source_counts.items(), 1):
        row = table.rows[i]
        row.cells[0].text = str(src)
        row.cells[1].text = str(cnt)
        row.cells[2].text = f"{cnt / len(df):.1%}"

    # Total row
    total_row = table.rows[-1]
    total_row.cells[0].text = "TOTAL"
    total_row.cells[1].text = str(len(df))
    total_row.cells[2].text = "100.0%"

    _format_table(table)

    doc.add_paragraph("")

    # Extraction metadata
    if engine and hasattr(engine, "pdf_count"):
        doc.add_paragraph(
            f"PST/OST emails scanned: {getattr(engine, 'email_count', 0)}",
            style=styles["BodyText"],
        )
        doc.add_paragraph(
            f"PDF attachments extracted: {getattr(engine, 'pdf_count', 0)}",
            style=styles["BodyText"],
        )

    doc.add_page_break()


def create_anomaly_detail_section(
    doc: Any,
    styles: Any,
    flags: list,
    df: pd.DataFrame,
    ctx: RenderContext | None = None,
) -> None:
    """Create detailed anomaly findings section."""
    if ctx is None:
        ctx = RenderContext()
    doc.add_paragraph(ctx.heading("detailed_findings"), style=styles["SectionHeader"])

    if not flags:
        doc.add_paragraph(
            "No specific anomalies were automatically detected. "
            "The timeline and statistical analysis sections may reveal patterns "
            "warranting manual investigation.",
            style=styles["BodyText"],
        )
    else:
        # Flag tuple schema from compute_dispute_flags:
        #   (type, date, amount, detail, severity)
        #   index:  0     1     2       3       4
        # The old code labelled index 0 as "Severity" (actually type),
        # index 2 as "Description" (actually amount), and index 3 as
        # "Records affected" (actually detail).  Fixed to use the correct
        # field names and the previously-missing severity at index 4.
        for i, flag in enumerate(flags, 1):
            doc.add_paragraph(
                f"Finding {i}: {flag[1] if len(flag) > 1 else 'Unknown'}",
                style=styles["SubSectionHeader"],
            )
            doc.add_paragraph(
                f"Type: {flag[0] if len(flag) > 0 else 'Unknown'}",
                style=styles["BodyText"],
            )
            doc.add_paragraph(
                f"Amount: {flag[2] if len(flag) > 2 else 'N/A'}",
                style=styles["BodyText"],
            )
            doc.add_paragraph(
                f"Detail: {flag[3] if len(flag) > 3 else 'No detail'}",
                style=styles["BodyText"],
            )
            doc.add_paragraph(
                f"Severity: {flag[4] if len(flag) > 4 else 'Unknown'}",
                style=styles["BodyText"],
            )
            doc.add_paragraph("")

    doc.add_page_break()


def create_timeline_section(
    doc: Any,
    styles: Any,
    df: pd.DataFrame,
    flags: list,
    ctx: RenderContext | None = None,
) -> None:
    """Create chronological timeline of events."""
    if ctx is None:
        ctx = RenderContext()
    doc.add_paragraph(ctx.heading("timeline"), style=styles["SectionHeader"])

    # Sort by date
    df_sorted = df.copy()
    df_sorted["_sort"] = df_sorted["Date"].apply(parse_to_sort_date)
    df_sorted = df_sorted.sort_values("_sort").reset_index(drop=True)

    table = doc.add_table(rows=min(len(df_sorted), 50) + 1, cols=5)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    headers = ["Date", "Source", "Entry Type", "Amount (£)", "Details"]
    for i, h in enumerate(headers):
        table.rows[0].cells[i].text = h

    for i, (_, row) in enumerate(df_sorted.head(50).iterrows()):
        r = table.rows[i + 1]
        r.cells[0].text = str(row.get("Date", "N/A"))
        r.cells[1].text = str(row.get("Source", "N/A"))
        r.cells[2].text = str(row.get("Entry Type", "N/A"))
        r.cells[3].text = fmt_money(row.get("Amount (£)", 0))
        r.cells[4].text = str(row.get("Details", ""))[:50]

    _format_table(table, font_size=7)

    if len(df_sorted) > 50:
        doc.add_paragraph(
            f"Showing first 50 of {len(df_sorted)} records. Full timeline in Appendix.",
            style=styles["Caption"],
        )

    doc.add_page_break()


def create_ofgem_comparison(
    doc: Any, styles: Any, df: pd.DataFrame, ctx: RenderContext | None = None
) -> None:
    """Create OFGEM price cap comparison section."""
    if ctx is None:
        ctx = RenderContext()
    doc.add_paragraph(ctx.heading("ofgem"), style=styles["SectionHeader"])

    doc.add_paragraph(
        "This section compares extracted unit rates against OFGEM price caps "
        "for the relevant periods. Note: OFGEM caps apply to standard variable "
        "tariffs; fixed tariffs may differ.",
        style=styles["BodyText"],
    )

    if "Unit Rate (p/kWh)" not in df.columns:
        doc.add_paragraph("No unit rate data available for comparison.", style=styles["BodyText"])
        doc.add_page_break()
        return

    df_rates = df.dropna(subset=["Unit Rate (p/kWh)"])
    df_rates = df_rates[df_rates["Unit Rate (p/kWh)"] != "N/A"]
    df_rates["rate_num"] = pd.to_numeric(df_rates["Unit Rate (p/kWh)"], errors="coerce")

    if df_rates["rate_num"].dropna().empty:
        doc.add_paragraph("No computable unit rates found.", style=styles["BodyText"])
        doc.add_page_break()
        return

    avg_rate = df_rates["rate_num"].mean()
    median_rate = df_rates["rate_num"].median()

    doc.add_paragraph(
        f"Average unit rate across all records: {fmt_number(avg_rate, 2)} p/kWh",
        style=styles["BodyText"],
    )
    doc.add_paragraph(
        f"Median unit rate: {fmt_number(median_rate, 2)} p/kWh", style=styles["BodyText"]
    )

    # OFGEM caps reference table — built dynamically from the shared
    # _load_ofgem_caps() data (edf_report.py) so the DOCX and PDF
    # generators always show the same values.  The old code hard-coded
    # a 7-row table that diverged from the PDF (e.g. "34.0" here vs
    # the correct 28.34 for Oct–Dec 2022).
    ofgem_caps = _load_ofgem_caps()
    # Only show recent/relevant periods (2022 onwards) to keep the
    # table readable; the full dataset goes back to 2019-Q1.
    recent_caps = {
        k: v for k, v in ofgem_caps.items() if int(k[:4]) >= 2022
    }
    # Human-readable quarter labels, e.g. "2022-Q4" → "Oct 2022 – Dec 2022"
    _q_start = {1: "Jan", 2: "Apr", 3: "Jul", 4: "Oct"}
    _q_end = {1: "Mar", 2: "Jun", 3: "Sep", 4: "Dec"}
    ofgem_rows = [("Period", "Electricity Cap (p/kWh)", "Source")]
    for q_key in sorted(recent_caps):
        year = int(q_key[:4])
        qn_ = int(q_key[-1])
        label = f"{_q_start[qn_]} {year} – {_q_end[qn_]} {year}"
        rate = recent_caps[q_key]["unit_rate"]
        ofgem_rows.append((label, f"{rate:.2f}", "OFGEM"))

    table = doc.add_table(rows=len(ofgem_rows), cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, row_data in enumerate(ofgem_rows):
        for j, val in enumerate(row_data):
            table.rows[i].cells[j].text = val

    _format_table(table, header_color="#10367A", font_size=9)

    doc.add_paragraph("")
    doc.add_paragraph(
        "Note: Price caps are for typical domestic consumption values. "
        "Actual rates vary by region, payment method, and tariff type.",
        style=styles["Caption"],
    )

    doc.add_page_break()


def create_statistical_analysis(
    doc: Any, styles: Any, df: pd.DataFrame, ctx: RenderContext | None = None
) -> None:
    """Create statistical analysis section."""
    if ctx is None:
        ctx = RenderContext()
    doc.add_paragraph(ctx.heading("statistical"), style=styles["SectionHeader"])

    if "Amount (£)" in df.columns:
        amounts = pd.to_numeric(df["Amount (£)"], errors="coerce").dropna()
        if not amounts.empty:
            doc.add_paragraph("Bill Amount Statistics:", style=styles["SubSectionHeader"])

            stats_data = [
                ("Metric", "Value"),
                ("Count", str(len(amounts))),
                ("Mean", fmt_money(amounts.mean())),
                ("Median", fmt_money(amounts.median())),
                ("Std Dev", fmt_money(amounts.std())),
                ("Min", fmt_money(amounts.min())),
                ("Max", fmt_money(amounts.max())),
                ("25th Percentile", fmt_money(amounts.quantile(0.25))),
                ("75th Percentile", fmt_money(amounts.quantile(0.75))),
            ]

            table = doc.add_table(rows=len(stats_data), cols=2)
            table.alignment = WD_TABLE_ALIGNMENT.CENTER
            for i, (label, val) in enumerate(stats_data):
                table.rows[i].cells[0].text = label
                table.rows[i].cells[1].text = val

            _format_table(table, header_color="#EBF3FA", font_size=10)

    doc.add_page_break()


def create_payment_analysis(
    doc: Any, styles: Any, df: pd.DataFrame, ctx: RenderContext | None = None
) -> None:
    """Create payment & credit analysis section."""
    if ctx is None:
        ctx = RenderContext()
    doc.add_paragraph(ctx.heading("payment"), style=styles["SectionHeader"])

    # Include both "Payment" and "Credit" entries — the PDF generator
    # (edf_report.py:1416) uses .isin(["Payment", "Credit"]) and the
    # DOCX side must match so credit notes appear in the totals.
    payments = df[df["Entry Type"].isin(["Payment", "Credit"])].copy()
    if not payments.empty and "Amount (£)" in payments.columns:
        pay_amounts = pd.to_numeric(payments["Amount (£)"], errors="coerce").dropna()
        if not pay_amounts.empty:
            doc.add_paragraph(f"Number of payments: {len(pay_amounts)}", style=styles["BodyText"])
            doc.add_paragraph(
                f"Total paid: {fmt_money(abs(pay_amounts.sum()))}", style=styles["BodyText"]
            )
            doc.add_paragraph(
                f"Average payment: {fmt_money(abs(pay_amounts.mean()))}", style=styles["BodyText"]
            )

            # Payment frequency
            if "Date" in payments.columns:
                pay_dates = payments["Date"].apply(parse_to_sort_date).dropna()
                if len(pay_dates) > 1:
                    diffs = pay_dates.sort_values().diff().dt.days.dropna()
                    if not diffs.empty:
                        doc.add_paragraph(
                            f"Average days between payments: {diffs.mean():.1f}",
                            style=styles["BodyText"],
                        )

    doc.add_page_break()


def create_forecast_section(
    doc: Any, styles: Any, df: pd.DataFrame, ctx: RenderContext | None = None
) -> None:
    """Create forecast & projection section."""
    if ctx is None:
        ctx = RenderContext()
    doc.add_paragraph(ctx.heading("forecast"), style=styles["SectionHeader"])

    doc.add_paragraph(
        "Based on historical billing patterns, the following projections are provided. "
        "These are estimates only and should not be relied upon for financial planning.",
        style=styles["BodyText"],
    )

    if "Amount (£)" in df.columns:
        amounts = pd.to_numeric(df["Amount (£)"], errors="coerce").dropna()
        if len(amounts) >= 3:
            recent_avg = amounts.tail(min(6, len(amounts))).mean()
            doc.add_paragraph(
                f"6-bill rolling average: {fmt_money(recent_avg)}", style=styles["BodyText"]
            )
            doc.add_paragraph(
                f"Estimated annual cost (12 bills): {fmt_money(recent_avg * 12)}",
                style=styles["BodyText"],
            )

    doc.add_page_break()


def create_data_quality_section(
    doc: Any, styles: Any, df: pd.DataFrame, ctx: RenderContext | None = None
) -> None:
    """Create data quality assessment section."""
    if ctx is None:
        ctx = RenderContext()
    doc.add_paragraph(ctx.heading("data_quality"), style=styles["SectionHeader"])

    total = len(df)
    missing = {}

    for col in ["Date", "Amount (£)", "Source", "Entry Type", "Invoice #", "Tariff", "Units (kWh)"]:
        if col in df.columns:
            n_missing = df[col].isna().sum() + (df[col] == "N/A").sum()
            missing[col] = n_missing

    table = doc.add_table(rows=len(missing) + 1, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    headers = ["Field", "Missing/NA", "Completeness"]
    for i, h in enumerate(headers):
        table.rows[0].cells[i].text = h

    for i, (col, n_miss) in enumerate(missing.items(), 1):
        row = table.rows[i]
        row.cells[0].text = col
        row.cells[1].text = str(n_miss)
        row.cells[2].text = f"{(total - n_miss) / total * 100:.1f}%"

    _format_table(table)

    doc.add_page_break()


def create_tariff_impact_section(
    doc: Any, styles: Any, df: pd.DataFrame, ctx: RenderContext | None = None
) -> None:
    """Create tariff impact analysis section."""
    if ctx is None:
        ctx = RenderContext()
    doc.add_paragraph(ctx.heading("tariff"), style=styles["SectionHeader"])

    if "Tariff" not in df.columns or df["Tariff"].isna().all() or (df["Tariff"] == "N/A").all():
        doc.add_paragraph(
            "No tariff data available in the extracted records. Tariff information is "
            "typically found on new-format (KI/KCR) invoices.",
            style=styles["BodyText"],
        )
        doc.add_page_break()
        return

    tariff_data = df.dropna(subset=["Tariff"])
    tariff_data = tariff_data[tariff_data["Tariff"] != "N/A"]

    if tariff_data.empty:
        doc.add_paragraph("No valid tariff records found.", style=styles["BodyText"])
        doc.add_page_break()
        return

    tariff_data = tariff_data.copy()
    tariff_data["unit_rate_num"] = pd.to_numeric(tariff_data["Unit Rate (p/kWh)"], errors="coerce")
    tariff_data = tariff_data.dropna(subset=["unit_rate_num"])

    if tariff_data.empty:
        doc.add_paragraph("No computable unit rates found.", style=styles["BodyText"])
        doc.add_page_break()
        return

    tariff_stats = (
        tariff_data.groupby("Tariff")
        .agg(
            count=("unit_rate_num", "count"),
            avg_rate=("unit_rate_num", "mean"),
            median_rate=("unit_rate_num", "median"),
            min_rate=("unit_rate_num", "min"),
            max_rate=("unit_rate_num", "max"),
        )
        .reset_index()
    )

    doc.add_paragraph("Unit Rate by Tariff:", style=styles["SubSectionHeader"])

    table = doc.add_table(rows=len(tariff_stats) + 1, cols=6)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    headers = ["Tariff", "Records", "Avg Rate (p/kWh)", "Median", "Min", "Max"]
    for i, h in enumerate(headers):
        table.rows[0].cells[i].text = h

    for i, (_, row) in enumerate(tariff_stats.iterrows(), 1):
        r = table.rows[i]
        r.cells[0].text = str(row["Tariff"])
        r.cells[1].text = str(int(row["count"]))
        r.cells[2].text = fmt_number(row["avg_rate"], 2)
        r.cells[3].text = fmt_number(row["median_rate"], 2)
        r.cells[4].text = fmt_number(row["min_rate"], 2)
        r.cells[5].text = fmt_number(row["max_rate"], 2)

    _format_table(table, font_size=8)

    doc.add_paragraph("")

    # Tariff changes
    tariff_data["_dt"] = tariff_data["Date"].apply(parse_to_sort_date)
    tariff_data = tariff_data.sort_values("_dt")
    changes = tariff_data["Tariff"].ne(tariff_data["Tariff"].shift()).cumsum()
    n_changes = int(changes.max()) if not changes.empty else 0

    doc.add_paragraph(f"Tariff Changes Detected: {n_changes}", style=styles["SubSectionHeader"])

    doc.add_page_break()


def create_appendix_methodology(
    doc: Any, styles: Any, config: Any, ctx: RenderContext | None = None
) -> None:
    """Create Methodology appendix."""
    if ctx is None:
        ctx = RenderContext()
    doc.add_paragraph(ctx.heading("appendix_methodology"), style=styles["SectionHeader"])

    doc.add_paragraph(
        "All billing records were extracted from three primary source types:",
        style=styles["BodyText"],
    )

    sources = [
        "• PDF Bills: EDF Energy invoices (both legacy and new KI/KCR formats) "
        "processed via pdfplumber with format-specific regex extraction.",
        "• HTM Export: EDF MyAccount 'Payments and Invoices' export parsed via "
        "BeautifulSoup with pattern matching for charge/payment/reversal entries.",
        "• PST/OST Email Archives: Outlook data files processed via libpff-python, "
        "extracting email bodies and PDF attachments.",
    ]

    for s in sources:
        doc.add_paragraph(s, style=styles["BodyText"])

    doc.add_paragraph("")

    doc.add_paragraph(
        "Amount extraction uses a tiered anchor-based approach with fallback to "
        "large-amount detection. Date extraction prioritizes bill/invoice date markers. "
        "Deduplication uses SHA-1 hashing of PDF content.",
        style=styles["BodyText"],
    )

    doc.add_paragraph("Configuration used:", style=styles["SubSectionHeader"])
    for key, val in config.items():
        doc.add_paragraph(f"• {key}: {val}", style=styles["BodyText"])


def create_appendix_glossary(doc: Any, styles: Any, ctx: RenderContext | None = None) -> None:
    """Create Glossary appendix."""
    if ctx is None:
        ctx = RenderContext()
    doc.add_paragraph(ctx.heading("appendix_glossary"), style=styles["SectionHeader"])
    glossary = {
        "KI-": "New-style EDF invoice reference prefix (e.g., KI-12345678)",
        "KCR-": "New-style EDF credit note reference prefix (e.g., KCR-87654321)",
        "A-": "EDF account number prefix (e.g., A-12345678)",
        "Standing Charge": "Daily fixed charge in pence per day, regardless of usage",
        "Unit Rate": "Price per kWh of electricity consumed (p/kWh)",
        "OFGEM Price Cap": "Regulatory maximum price per kWh for standard variable tariffs",
        "Period Charge": "Total charges for a specific billing period",
        "Current Balance": "Running account balance including all historical charges/payments",
        "Estimated Reading": "Meter reading estimated by supplier, not actual",
        "Actual Reading": "Meter reading provided by customer or smart meter",
    }

    table = doc.add_table(rows=len(glossary) + 1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    for i, (term, definition) in enumerate(glossary.items(), 1):
        table.rows[i].cells[0].text = term
        table.rows[i].cells[1].text = definition

    _format_table(table, header_color="#EBF3FA", font_size=9)

    doc.add_page_break()


def create_appendix_full_evidence(
    doc: Any,
    styles: Any,
    df: pd.DataFrame,
    filtered: Any = None,
    ctx: RenderContext | None = None,
) -> None:
    """Create Full Evidence Table appendix, plus an optional Filtered Records sub-table."""
    if ctx is None:
        ctx = RenderContext()
    doc.add_paragraph(ctx.heading("appendix_full_evidence"), style=styles["SectionHeader"])

    doc.add_paragraph(
        "This appendix contains the complete set of billing records used in this analysis. "
        "Records are sorted chronologically by date.",
        style=styles["BodyText"],
    )

    if df.empty:
        doc.add_paragraph("No records available.", style=styles["BodyText"])
        doc.add_page_break()
        return

    # Ensure _dt column for sorting
    if "_dt" not in df.columns:
        df["_dt"] = df["Date"].apply(parse_to_sort_date)
    df_sorted = df.sort_values("_dt").reset_index(drop=True)

    # Table header
    evidence_header = [
        "Date",
        "Source",
        "Entry Type",
        "Amount (£)",
        "Period Charge (£)",
        "Period From",
        "Period To",
        "Invoice #",
        "Reading",
        "Units (kWh)",
        "Standing Chg (p/day)",
        "Attachment",
        "Details",
    ]

    evidence_data = [evidence_header]

    for _, row in df_sorted.iterrows():
        evidence_data.append(
            [
                fmt_date(row.get("Date", "")),
                str(row.get("Source", ""))[:30],
                str(row.get("Entry Type", "")),
                fmt_money(row.get("Amount (£)", 0)),
                fmt_money(row.get("Period Charge (£)", 0))
                if row.get("Period Charge (£)") not in ("", "N/A", None)
                else "N/A",
                fmt_date(row.get("Period From", "")),
                fmt_date(row.get("Period To", "")),
                str(row.get("Invoice #", ""))[:15],
                str(row.get("Reading", ""))[:15],
                str(row.get("Units (kWh)", ""))[:10],
                str(row.get("Standing Chg (p/day)", ""))[:10],
                str(row.get("Attachment Name", ""))[:20],
                str(row.get("Details", ""))[:50],
            ]
        )

    table = doc.add_table(rows=len(evidence_data), cols=len(evidence_header))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"

    # Set header row
    for j, header in enumerate(evidence_header):
        table.rows[0].cells[j].text = header
        for paragraph in table.rows[0].cells[j].paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
                run.font.color.rgb = WHITE
        table.rows[0].cells[j]._element.get_or_add_tcPr().append(
            parse_xml(
                '<w:shd {} w:fill="10367A"/>'.format(
                    'xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"'
                )
            )
        )

    # Fill data rows
    for i, row_data in enumerate(evidence_data[1:], 1):
        for j, cell_text in enumerate(row_data):
            table.rows[i].cells[j].text = str(cell_text)
            if i % 2 == 0:
                table.rows[i].cells[j]._element.get_or_add_tcPr().append(
                    parse_xml(
                        '<w:shd {} w:fill="EBF3FA"/>'.format(
                            'xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"'
                        )
                    )
                )

    doc.add_page_break()

    # Add filtered records if provided
    if filtered:
        cont_label = ctx.short_label("appendix_full_evidence").rstrip(".")
        cont_heading = (
            f"{cont_label}. (cont.) Filtered Records (Below £500 Threshold)"
            if cont_label
            else "Filtered Records (Below £500 Threshold)"
        )
        doc.add_paragraph(cont_heading, style=styles["SectionHeader"])

        filt_data = [
            [
                "Date",
                "Source",
                "Entry Type",
                "Amount (£)",
                "Period Charge (£)",
                "Period From",
                "Period To",
                "Invoice #",
                "Reading",
                "Units (kWh)",
                "Standing Chg (p/day)",
                "Attachment",
                "Details",
            ]
        ]
        for row in filtered:
            filt_data.append(
                [
                    fmt_date(row.get("Date", "")),
                    str(row.get("Source", ""))[:30],
                    str(row.get("Entry Type", "")),
                    fmt_money(row.get("Amount (£)", 0)),
                    fmt_money(row.get("Period Charge (£)", 0))
                    if row.get("Period Charge (£)") not in ("", "N/A", None)
                    else "N/A",
                    fmt_date(row.get("Period From", "")),
                    fmt_date(row.get("Period To", "")),
                    str(row.get("Invoice #", ""))[:15],
                    str(row.get("Reading", ""))[:15],
                    str(row.get("Units (kWh)", ""))[:10],
                    str(row.get("Standing Chg (p/day)", ""))[:10],
                    str(row.get("Attachment Name", ""))[:20],
                    str(row.get("Details", ""))[:50],
                ]
            )

        if len(filt_data) > 1:
            table2 = doc.add_table(rows=len(filt_data), cols=len(filt_data[0]))
            table2.alignment = WD_TABLE_ALIGNMENT.CENTER
            table2.style = "Table Grid"

            for j, header in enumerate(filt_data[0]):
                table2.rows[0].cells[j].text = header
                for paragraph in table2.rows[0].cells[j].paragraphs:
                    for run in paragraph.runs:
                        run.font.bold = True
                        run.font.color.rgb = WHITE
                table2.rows[0].cells[j]._element.get_or_add_tcPr().append(
                    parse_xml(
                        '<w:shd {} w:fill="FFA500"/>'.format(
                            'xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"'
                        )
                    )
                )

            for i, row_data in enumerate(filt_data[1:], 1):
                for j, cell_text in enumerate(row_data):
                    table2.rows[i].cells[j].text = str(cell_text)
                    if i % 2 == 0:
                        table2.rows[i].cells[j]._element.get_or_add_tcPr().append(
                            parse_xml(
                                '<w:shd {} w:fill="EBF3FA"/>'.format(
                                    'xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"'
                                )
                            )
                        )

        doc.add_page_break()


# =============================================================================
# MAIN GENERATOR
# =============================================================================


def generate_ombudsman_docx(
    records: list[dict],
    output_path: str,
    config: dict,
    engine: Any,
    filtered: list | None = None,
) -> str:
    """
    Generate a professional DOCX report for Energy Ombudsman submission.

    Args:
        records: List of extracted billing records
        output_path: Path to save the DOCX
        config: Configuration dictionary
        engine: EvidenceEngine instance (for metadata)
        filtered: Filtered-out records (below threshold)

    Returns:
        Path to generated DOCX
    """
    if not records:
        raise ValueError("No records to report on")

    df = pd.DataFrame(records)
    if df.empty:
        raise ValueError("Records DataFrame is empty")

    # Ensure required columns exist with defaults
    required_cols = {
        "Date": "01/01/1970",
        "Source": "Unknown",
        "Amount (£)": 0.0,
        "Period From": "N/A",
        "Period To": "N/A",
        "Invoice #": "N/A",
        "Period Charge (£)": "N/A",
        "Entry Type": "Unknown",
        "Reading": "N/A",
        "Units (kWh)": "N/A",
        "Standing Chg (p/day)": "N/A",
        "Attachment Name": "N/A",
        "Details": "",
        "Logic Used": "",
    }
    for col, default in required_cols.items():
        if col not in df.columns:
            df[col] = default

    df["_sort"] = df["Date"].apply(parse_to_sort_date)
    df = df.sort_values("_sort").reset_index(drop=True)

    # Account reference
    acc_ref = config.get("report_account_ref") or config.get("acc_num") or "Unknown"

    # Period bounds
    dates_parsed = df["Date"].apply(parse_to_sort_date)
    valid_dates = dates_parsed.dropna()
    period_start = fmt_date(valid_dates.min()) if not valid_dates.empty else "Unknown"
    period_end = fmt_date(valid_dates.max()) if not valid_dates.empty else "Unknown"

    # Financial totals
    charges = df[df["Amount (£)"] > 0]["Amount (£)"].astype(float).sum()
    payments = df[df["Amount (£)"] < 0]["Amount (£)"].astype(float).sum() * -1

    # Compute dispute flags from the data
    from edf_collector import compute_dispute_flags

    # Ensure _dt column exists for the analysis
    if "_dt" not in df.columns:
        df["_dt"] = df["Date"].apply(parse_to_sort_date)
    df_sorted = df.sort_values("_dt").reset_index(drop=True)

    # Mean daily rate — shared logic (imported from edf_report).
    mean_daily = _compute_mean_daily(df_sorted)

    flags, flag_counts = compute_dispute_flags(df_sorted, mean_daily)

    # Section selection: only include sections in config["report_sections"]
    enabled_sections = set(config.get("report_sections", []))
    # Backward compatibility: if not specified, enable all
    all_sections = {
        "cover",
        "toc",
        "exec_summary",
        "key_findings",
        "evidence_index",
        "detailed_findings",
        "timeline",
        "ofgem",
        "statistical",
        "payment",
        "forecast",
        "data_quality",
        "tariff",
        "appendix_methodology",
        "appendix_glossary",
        "appendix_full_evidence",
    }
    if not enabled_sections:
        enabled_sections = all_sections

    def section_enabled(key: str) -> bool:
        return key in enabled_sections

    # RenderContext derives every section's number/letter label from the
    # registry and the user's selection. The TOC and every section body
    # consume the same context so headline numbers and headings match.
    render_ctx = RenderContext(enabled_sections)

    # Build document
    doc = Document()

    # Set default margins
    for section in doc.sections:
        section.top_margin = Cm(MARGIN_CM)
        section.bottom_margin = Cm(MARGIN_CM)
        section.left_margin = Cm(MARGIN_CM)
        section.right_margin = Cm(MARGIN_CM)

    # Create styles
    styles = _get_or_create_styles(doc)

    # Add footer
    _add_footer(doc)

    # === COVER PAGE ===
    if section_enabled("cover"):
        create_cover_page(
            doc, styles, acc_ref, period_start, period_end, datetime.now().strftime("%d %B %Y")
        )

    # === TABLE OF CONTENTS ===
    if section_enabled("toc"):
        create_table_of_contents(doc, styles, render_ctx)

    # === SECTION DISPATCH (data-driven — keys/ordering live in REPORT_SECTIONS) ===
    section_builders: dict[str, tuple] = {
        "exec_summary": (
            lambda: {
                "doc": doc,
                "styles": styles,
                "df": df,
                "config": config,
                "acc_ref": acc_ref,
                "flag_counts": flag_counts,
                "n_records": len(records),
                "charges": charges,
                "payments": payments,
                "period_start": period_start,
                "period_end": period_end,
            },
            lambda kwargs: create_executive_summary(**kwargs),
        ),
        "key_findings": (
            lambda: {"doc": doc, "styles": styles, "flags": flags},
            lambda kwargs: create_key_findings_table(**kwargs),
        ),
        "evidence_index": (
            lambda: {"doc": doc, "styles": styles, "df": df, "engine": engine},
            lambda kwargs: create_evidence_index(**kwargs),
        ),
        "detailed_findings": (
            lambda: {"doc": doc, "styles": styles, "flags": flags, "df": df},
            lambda kwargs: create_anomaly_detail_section(**kwargs),
        ),
        "timeline": (
            lambda: {"doc": doc, "styles": styles, "df": df, "flags": flags},
            lambda kwargs: create_timeline_section(**kwargs),
        ),
        "ofgem": (
            lambda: {"doc": doc, "styles": styles, "df": df},
            lambda kwargs: create_ofgem_comparison(**kwargs),
        ),
        "statistical": (
            lambda: {"doc": doc, "styles": styles, "df": df},
            lambda kwargs: create_statistical_analysis(**kwargs),
        ),
        "payment": (
            lambda: {"doc": doc, "styles": styles, "df": df},
            lambda kwargs: create_payment_analysis(**kwargs),
        ),
        "forecast": (
            lambda: {"doc": doc, "styles": styles, "df": df},
            lambda kwargs: create_forecast_section(**kwargs),
        ),
        "data_quality": (
            lambda: {"doc": doc, "styles": styles, "df": df},
            lambda kwargs: create_data_quality_section(**kwargs),
        ),
        "tariff": (
            lambda: {"doc": doc, "styles": styles, "df": df},
            lambda kwargs: create_tariff_impact_section(**kwargs),
        ),
        "appendix_methodology": (
            lambda: {"doc": doc, "styles": styles, "config": config},
            lambda kwargs: create_appendix_methodology(**kwargs),
        ),
        "appendix_glossary": (
            lambda: {"doc": doc, "styles": styles},
            lambda kwargs: create_appendix_glossary(**kwargs),
        ),
        "appendix_full_evidence": (
            lambda: {"doc": doc, "styles": styles, "df": df, "filtered": filtered},
            lambda kwargs: create_appendix_full_evidence(**kwargs),
        ),
    }

    for spec in REPORT_SECTIONS:
        if not section_enabled(spec.key):
            continue
        entry = section_builders.get(spec.key)
        if entry is None:
            raise RuntimeError(
                f"REPORT_SECTIONS lists '{spec.key}' but no DOCX builder is wired "
                f"in generate_ombudsman_docx. Add it to section_builders."
            )
        arg_factory, invoke = entry
        kwargs = arg_factory()
        kwargs["ctx"] = render_ctx
        invoke(kwargs)

    # Save
    doc.save(output_path)
    return output_path


def generate_docx_from_gui(records, output_path, config, engine, filtered=None):
    """Wrapper for GUI integration."""
    try:
        path = generate_ombudsman_docx(records, output_path, config, engine, filtered)
        return True, f"Professional DOCX report generated:\n{path}"
    except Exception as e:
        return False, f"Failed to generate DOCX:\n{e}"
