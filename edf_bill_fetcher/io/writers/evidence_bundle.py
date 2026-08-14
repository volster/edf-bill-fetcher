"""Save referenced evidence files and build a themed DOCX bundle index.

Canonical home: ``edf_bill_fetcher.io.writers.evidence_bundle`` (absorbed
from the top-level ``evidence_bundle.py`` during the shim-removal wave).

This module implements Stream P5 of the SAP/Evidence Bundle design
(``scratch/superpowers/specs/2026-07-16-sap-dumps-and-evidence-bundle-design.md``):

- ``save_evidence_files(expr_df, source_paths, dest_dir)`` copies into a flat
  ``dest_dir`` every file the Evidence Report references via its
  ``Attachment Name`` column. Collisions are handled by appending ``-2``,
  ``-3`` etc.
- ``build_bundle_index(evidence_df, saved_files, output_path)`` writes a
  themed-section DOCX bundle index using python-docx. The section-to-row
  mapping follows a 3-layer inference (prefix / content fingerprint / default
  with the ``⚠  unprefixed`` flag).
"""

from __future__ import annotations

import os
import re
import shutil
from collections import defaultdict
from collections.abc import Callable
from typing import Any

import pandas as pd

# ---------------------------------------------------------------------------
# Section inference (Stream P5, spec §7.3)
# ---------------------------------------------------------------------------

# Section labels -- order matters when we render the bundle.
SECTION_LABELS: dict[str, str] = {
    "A": "A — Ombudsman and Regulatory",
    "B": "B — Correspondence",
    "C": "C — PSR",
    "D": "D — Invoices",
    "E": "E — Meter Readings",
    "F": "F — Balance",
    "G": "G — Calls",
}

# Filename-prefix regex: ``A1 - foo.pdf`` or ``B12 - bar.pdf``. The captured
# group is the upper-case letter identifying the section.
_PREFIX_RE = re.compile(r"^([A-G])\d*\s*-\s*", re.IGNORECASE)

# Content-fingerprint regexes (Stream P5, spec §7.3 layer 2): recognisable
# EDF dump filenames without prefix.
_METER_HISTORY_RE = re.compile(r"Meter-?Read-?History", re.IGNORECASE)
_FINANCIAL_TX_RE = re.compile(r"Financial-?Transactions", re.IGNORECASE)
_CONTRACT_HISTORY_RE = re.compile(r"Contract-?and-?Product-?Change-?History", re.IGNORECASE)


def _infer_section(attachment_name: str, fallback: str = "D") -> tuple[str, bool]:
    """Return ``(section_letter, is_unprefixed)`` for a single row.

    Three-layer inference:
      1. Filename suffix-letter prefix (``A1 -``, ``B12 -``, ...).
      2. Content fingerprint (``[[Meter-Read-History`` → E, etc.).
      3. Default ``D — Invoices`` -- the unrouted-but-categorised catchall.
         Unprefixed files are flagged so the bundle can display a
         ``⚠  unprefixed`` marker.
    """
    name = str(attachment_name or "")
    base = os.path.basename(name)
    m = _PREFIX_RE.match(base)
    if m:
        letter = m.group(1).upper()
        if letter in SECTION_LABELS:
            return letter, False
    # Layer 2: fingerprint by filename + context.
    if _METER_HISTORY_RE.search(base):
        # Meter-Read-History has no £/kWh readings per-row -- route to E.
        return "E", False
    if _FINANCIAL_TX_RE.search(base):
        return "F", False
    if _CONTRACT_HISTORY_RE.search(base):
        return "D", False
    # Layer 3: default-to-D, flagged as unprefixed.
    return fallback, True


# ---------------------------------------------------------------------------
# Save referenced files into a flat dest dir
# ---------------------------------------------------------------------------

_WINDOWS_ILLEGAL_RE = re.compile(r'[<>:"/\\|?*]')
_WS_RUN_RE = re.compile(r"\s+")


def sanitise_filename(name: str) -> str:
    """Return a filesystem-safe basename (no dirs, no illegal chars)."""
    base = str(name or "").strip()
    base = _WINDOWS_ILLEGAL_RE.sub("_", base)
    base = _WS_RUN_RE.sub("_", base).strip(" .")
    return base or "attachment"


def save_evidence_files(
    evidence_df: pd.DataFrame,
    source_paths: dict[str, str],
    dest_dir: str,
    log: Callable[[str], Any] | None = None,
) -> dict[str, str]:
    """Copy every referred file into ``dest_dir``.

    ``source_paths[attachment_name] -> absolute_path`` is the reverse lookup
    the GUI / CLI builds while walking the local PDF / PST / HTM corpus. Any
    missing path is logged and skipped. Collisions (``Attachment_N.pdf`` in
    multiple PST emails, for example) are deduped by suffixing ``-2``, ``-3``
    etc. on the destination basename.

    Returns a dict ``{attachment_name: destination_path}`` for every file
    successfully copied.
    """
    if log is None:

        def log(_msg: str) -> None:
            pass

    os.makedirs(dest_dir, exist_ok=True)
    saved: dict[str, str] = {}
    used_names: set[str] = set()

    if evidence_df is None or "Attachment Name" not in evidence_df.columns:
        return saved

    invoice_to_att: dict[str, str] = {}
    multi_att_invoices: set[str] = set()
    if "Invoice #" in evidence_df.columns:
        for _, r in evidence_df.iterrows():
            inv_num = str(r.get("Invoice #", "")).strip()
            att = str(r.get("Attachment Name", "N/A"))
            if not inv_num or inv_num in ("N/A", "None", "nan"):
                continue
            if att in ("N/A", "", "None", "nan"):
                continue
            if inv_num in invoice_to_att and invoice_to_att[inv_num] != att:
                multi_att_invoices.add(inv_num)
            else:
                invoice_to_att.setdefault(inv_num, att)
        for inv_num in multi_att_invoices:
            invoice_to_att.pop(inv_num, None)

    for _, r in evidence_df.iterrows():
        att = str(r.get("Attachment Name", "N/A"))
        if att in ("N/A", "", "None", "nan"):
            continue
        if att in saved:
            continue
        src = source_paths.get(att)
        if not src or not os.path.exists(src):
            log(f"evidence_files: missing source for {att!r}")
            continue
        dest_base = att
        inv_num = str(r.get("Invoice #", "")).strip()
        if inv_num and inv_num not in ("N/A", "None", "nan") and invoice_to_att.get(inv_num) == att:
            dest_base = sanitise_filename(inv_num) + os.path.splitext(att)[1]
        n = 2
        while dest_base in used_names:
            stem, ext = os.path.splitext(att)
            dest_base = f"{stem}-{n}{ext}"
            n += 1
        used_names.add(dest_base)
        dest_path = os.path.join(dest_dir, dest_base)
        shutil.copy2(src, dest_path)
        saved[att] = dest_path
    return saved


# ---------------------------------------------------------------------------
# Build the themed-section DOCX bundle index
# ---------------------------------------------------------------------------


def build_bundle_index(
    evidence_df: pd.DataFrame,
    saved_files: dict[str, str],
    output_path: str,
    account: str = "",
) -> None:
    """Generate the themed-section DOCX bundle index.

    The doc has one ``<h2>`` per section (A..G, in the order the spec
    defines) and one bullet per saved file. Rows are routed via
    :func:`_infer_section`. Unprefixed rows are tagged with the
    ``⚠  unprefixed`` marker on the bullet line.
    """
    import docx  # local import so unit tests can mock it out cheaply.

    doc = docx.Document()
    h = doc.add_heading("EDF Evidence Bundle Index", level=1)
    if account:
        h.add_run(f"  |  Account {account}")
    doc.add_paragraph(
        "Every file referenced by the Evidence Report is reproduced in the "
        "`evidence_files/` subfolder. This index lists them grouped by the "
        "section letter assigned during inference."
    )

    # Bucket rows by section letter.
    by_section: dict[str, list[tuple[str, str, bool]]] = defaultdict(list)
    if evidence_df is not None and not evidence_df.empty:
        for _, r in evidence_df.iterrows():
            att = str(r.get("Attachment Name", "N/A"))
            if att in ("N/A", "", "None", "nan"):
                continue
            if att not in saved_files:
                continue
            section_letter, is_unprefixed = _infer_section(att)
            date = str(r.get("Date", "") or "")
            by_section[section_letter].append((att, date, is_unprefixed))

    # Emit one section heading + bullet list per letter present.
    for letter, label in SECTION_LABELS.items():
        rows = by_section.get(letter, [])
        if not rows:
            continue
        doc.add_heading(label, level=2)
        # Render one bullet per saved file.
        for att, date, is_unprefixed in rows:
            saved_path = saved_files.get(att, att)
            bullet_text = f"{os.path.basename(saved_path)}"
            if date:
                bullet_text += f"  —  date {date}"
            if is_unprefixed:
                bullet_text += "  ⚠  unprefixed"
            doc.add_paragraph(bullet_text, style="List Bullet")

    doc.save(output_path)
